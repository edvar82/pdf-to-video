"""Parser para o arquivo script.docx, extraindo tokens de slides, pausas e vignette."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document

from .types import PauseToken, ScriptToken, SlideToken, VignetteToken

TOKEN_PATTERN = re.compile(r"\[(slide_(\d+)|short_pause|long_pause|vignette)\]")


def _iter_tokens_from_text(text: str, short_pause_seconds: float, long_pause_seconds: float) -> Iterable[ScriptToken]:
    """Gera tokens na ordem em que aparecem no texto plano do script.
    
    Remove tags de pause e vignette que aparecem dentro de parágrafos e mantém apenas
    as ocorrências que estão no início de linhas ou após um token de slide.
    """
    # Primeiro, dividir o texto em linhas para processar corretamente
    lines = text.split('\n')
    processed_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Se a linha contém apenas tokens (slide, pause ou vignette), mantém como está
        if re.fullmatch(r'(\[(slide_\d+|short_pause|long_pause|vignette)\]\s*)+', line):
            processed_text.append(line)
        else:
            # Remove tokens de pause e vignette que aparecem no meio do texto
            # mas mantém tokens de slide
            cleaned_line = re.sub(r'\s*\[(short_pause|long_pause|vignette)\]\s*', ' ', line)
            cleaned_line = cleaned_line.strip()
            if cleaned_line:
                processed_text.append(cleaned_line)
    
    # Agora processar os tokens do texto limpo
    cleaned_text = '\n'.join(processed_text)
    
    for match in TOKEN_PATTERN.finditer(cleaned_text):
        raw = match.group(1)
        if raw.startswith("slide_"):
            index_text = match.group(2)
            if index_text is not None:
                yield SlideToken(slide_index=int(index_text))
        elif raw == "short_pause":
            yield PauseToken(seconds=short_pause_seconds)
        elif raw == "long_pause":
            yield PauseToken(seconds=long_pause_seconds)
        elif raw == "vignette":
            yield VignetteToken()


def parse_script_docx(docx_path: Path, short_pause_seconds: float, long_pause_seconds: float) -> List[ScriptToken]:
    """Lê o script.docx e retorna a lista ordenada de tokens reconhecidos.

    Caso o arquivo esteja vazio ou sem tokens, retorna uma lista vazia.
    """

    if not docx_path.exists():
        return []
    document = Document(str(docx_path))
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    joined = "\n".join(full_text)
    tokens = list(_iter_tokens_from_text(joined, short_pause_seconds, long_pause_seconds))
    return tokens


def extract_slide_texts(docx_path: Path) -> Dict[int, str]:
    """Extrai textos por slide com base nos marcadores [slide_XX] no script.

    Retorna um dicionário {index: texto}. Linhas entre um [slide_XX] e o próximo marcador
    são agregadas como texto do slide. Remove tags de pause e vignette do texto.
    """

    texts: Dict[int, List[str]] = {}
    if not docx_path.exists():
        return {}
    document = Document(str(docx_path))
    current_idx: int = -1
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if not line:
            continue
        m = re.search(r"\[slide_(\d+)\]", line)
        if m:
            current_idx = int(m.group(1))
            texts.setdefault(current_idx, [])
            # Remove a tag de slide da linha e adiciona o restante se houver
            remaining = re.sub(r'\[slide_\d+\]', '', line).strip()
            # Remove tags de pause e vignette
            remaining = re.sub(r'\s*\[(short_pause|long_pause|vignette)\]\s*', ' ', remaining).strip()
            if remaining:
                texts[current_idx].append(remaining)
            continue
        if current_idx != -1:
            # Remove tags de pause e vignette do texto
            cleaned_line = re.sub(r'\s*\[(short_pause|long_pause|vignette)\]\s*', ' ', line).strip()
            if cleaned_line:
                texts.setdefault(current_idx, []).append(cleaned_line)
    return {k: " ".join(v).strip() for k, v in texts.items()}


